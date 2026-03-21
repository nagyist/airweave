"""Direct text extraction from DOCX files using python-docx."""

from __future__ import annotations

import asyncio
import os
from typing import Any, Optional

from airweave.core.logging import logger
from airweave.domains.sync_pipeline.exceptions import SyncFailureError

MIN_TOTAL_CHARS = 50

_HEADING_MAP = (
    ("heading 1", "# "),
    ("heading 2", "## "),
    ("heading 3", "### "),
    ("heading", "#### "),
)


def _format_paragraph(para: Any) -> Optional[str]:
    text = para.text.strip()
    if not text:
        return None

    style_name = (para.style.name or "").lower() if para.style else ""

    for keyword, prefix in _HEADING_MAP:
        if keyword in style_name:
            return f"{prefix}{text}"

    if "list" in style_name:
        return f"- {text}"

    return text


def _format_table(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) > 1:
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)

    return "\n".join(rows)


async def extract_docx_text(path: str) -> Optional[str]:
    """Extract text from a DOCX and return markdown."""
    try:
        from docx import Document
    except ImportError:
        raise SyncFailureError("python-docx required for DOCX text extraction but not installed")

    def _extract() -> Optional[str]:
        name = os.path.basename(path)

        try:
            doc = Document(path)
        except Exception as exc:
            logger.warning(f"Failed to open DOCX {name}: {exc}")
            return None

        parts: list[str] = []

        for para in doc.paragraphs:
            line = _format_paragraph(para)
            if line:
                parts.append(line)

        for table in doc.tables:
            md_table = _format_table(table)
            if md_table:
                parts.append(md_table)

        markdown = "\n\n".join(parts)

        total_chars = len(markdown.strip())
        if total_chars < MIN_TOTAL_CHARS:
            logger.debug(f"DOCX {name}: only {total_chars} chars extracted, insufficient")
            return None

        logger.debug(f"DOCX {name}: extracted {total_chars} chars")
        return markdown

    return await asyncio.to_thread(_extract)
